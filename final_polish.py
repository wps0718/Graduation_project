# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

SRC = r'D:\Desktop\毕业设计(论文）最最最总_修正版.docx'
OUT = r'D:\Desktop\毕业设计(论文）最最最总_最终版.docx'
IMG_DIR = r'D:\Desktop\笔记\使用照片'

doc = Document(SRC)

# ============================================================
# Part 1: Fix acknowledgments (P251) and reference list (P255, P257)
# ============================================================

# Fix P251 - acknowledgments
p251 = doc.paragraphs[251]
old_text = p251.text
new_ack = '在本毕业设计完成之际，衷心感谢指导老师在选题确定、方案设计和论文撰写全过程中给予的耐心指导与宝贵建议，为项目的顺利完成提供了重要保障。感谢各位授课老师在大学期间的悉心教导，为我打下了扎实的专业基础。感谢在广轻大遇到的每一位老师和同学，你们的陪伴与帮助让我的大学时光充实而难忘。同时感谢家人一直以来的理解与鼓励，他们始终是我最坚强的后盾。最后，愿我们都能德能兼备，学以成之。'
for run in p251.runs:
    run.text = ''
p251.runs[0].text = new_ack
print('P251: Acknowledgments updated')

# Fix P255 - add [1] prefix
p255 = doc.paragraphs[255]
old = p255.text
if not old.startswith('[1]'):
    new_text = '[1] ' + old.lstrip()
    for run in p255.runs:
        run.text = ''
    p255.runs[0].text = new_text
    print(f'P255: Added [1] prefix')

# Fix P257 - split [3] and [7] onto separate lines
p257_text = doc.paragraphs[257].text
# Expected: "[3] MyBatis-Plus官方团队. MyBatis-Plus官方文档[EB/OL]. https://baomidou.com, 2024.[7] Carlson J. Redis in Action[M]. Manning Publications, 2016."
if '[7]' in p257_text and '[3]' in p257_text:
    # Split into [3] only - [7] will be handled separately
    part3 = '[3] MyBatis-Plus官方团队. MyBatis-Plus官方文档[EB/OL]. https://baomidou.com, 2024.'
    for run in doc.paragraphs[257].runs:
        run.text = ''
    doc.paragraphs[257].runs[0].text = part3
    print('P257: Split [3] and [7] - [3] kept in P257')

# Check if [7] Redis needs to be added as a separate entry
# Search for existing [9] or Carlson in nearby paragraphs
redis_found = False
for i in range(260, 268):
    if i < len(doc.paragraphs):
        if 'Carlson' in doc.paragraphs[i].text or 'Redis in Action' in doc.paragraphs[i].text:
            redis_found = True
            print(f'P{i}: Redis reference already exists')
            break

if not redis_found:
    # Find a blank line after P262 to insert [9]
    for i in range(263, min(270, len(doc.paragraphs))):
        if doc.paragraphs[i].text.strip() == '':
            doc.paragraphs[i].add_run('[9] Carlson J. Redis in Action[M]. Manning Publications, 2016.')
            print(f'P{i}: Added [9] Redis in Action')
            redis_found = True
            break

if not redis_found:
    print('WARNING: Could not place [9] Redis reference')

# ============================================================
# Part 2: Insert 6 images with captions
# Insert from bottom to top to avoid index shifting
# ============================================================

insertions = [
    # (after_paragraph_index, image_filename, caption, figure_number)
    (220, '订单状态机流转图.png',       '图5-3  订单状态机流转图'),
    (216, 'WebSocket 即时通讯架构图.png', '图5-2  WebSocket即时通讯架构图'),
    (179, '数据库简化 ER 关系图.png',   '图4-2  数据库ER关系图'),
    (174, '技术架构分层图.png',         '图4-1  系统技术架构分层图'),
    (169, '商品发布审核流程图.png',     '图3-2  商品发布审核流程图'),
    (165, '校园认证流程图.png',         '图3-1  校园认证流程图'),
]

def insert_image_with_caption(doc, after_idx, img_path, caption_text):
    """Insert image + caption after a specific paragraph using XML manipulation."""
    # Add paragraphs at the end first
    p_blank1 = doc.add_paragraph()
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(14))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.size = Pt(10.5)
    run_cap.font.name = '宋体'
    run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    p_blank2 = doc.add_paragraph()

    # Move these 4 paragraphs from end to after target paragraph
    target = doc.paragraphs[after_idx]._element
    body = target.getparent()
    all_paras = body.findall(qn('w:p'))

    # Last 4 elements are our new ones: blank1, img, cap, blank2
    new_elems = [all_paras[-4], all_paras[-3], all_paras[-2], all_paras[-1]]

    # Remove from end and insert after target
    for elem in new_elems:
        body.remove(elem)

    insert_after = target
    for elem in new_elems:
        insert_after.addnext(elem)
        insert_after = elem

# Sort by index descending so bottom inserts first
insertions_sorted = sorted(insertions, key=lambda x: x[0], reverse=True)

for after_idx, img_file, caption in insertions_sorted:
    img_path = f'{IMG_DIR}\\{img_file}'
    print(f'Inserting {img_file} after P{after_idx}...')
    insert_image_with_caption(doc, after_idx, img_path, caption)

# ============================================================
# Save
# ============================================================
doc.save(OUT)
print(f'\nSaved to: {OUT}')

# Verify
import zipfile, re
total = 0
chinese = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        total += len(text)
        chinese += len(re.findall(r'[一-鿿]', text))

with zipfile.ZipFile(OUT) as z:
    imgs = [f for f in z.namelist() if f.startswith('word/media/')]
    print(f'Images: {len(imgs)}')
    print(f'Total chars: {total}, Chinese chars: {chinese}')
