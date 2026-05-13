# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
import copy

SRC = r'D:\Desktop\毕业设计(论文）最最最总_修正版.docx'
OUT = r'D:\Desktop\毕业设计(论文）最最最总_最终版.docx'

doc = Document(SRC)

# Step 1: Remove incorrect outlineLvl from P113 (body text of 1.3)
p113 = doc.paragraphs[113]
pPr = p113._element.find(qn('w:pPr'))
if pPr is not None:
    outlineLvl = pPr.find(qn('w:outlineLvl'))
    if outlineLvl is not None:
        pPr.remove(outlineLvl)
        print('P113: Removed incorrect outlineLvl from body text')

# Step 2: Set outlineLvl for all chapter/section/subsection headings
def set_outline_level(para, level):
    """Set outline level on a paragraph (0=chapter, 1=section, 2=subsection)."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = para._element.makeelement(qn('w:pPr'), {})
        para._element.insert(0, pPr)
    # Remove existing outlineLvl if any
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    # Add new outlineLvl
    ol = para._element.makeelement(qn('w:outlineLvl'), {qn('w:val'): str(level)})
    pPr.append(ol)

# Map: paragraph_index -> outline_level
heading_fixes = {
    # Chapter 2
    130: 0,  # 二、系统相关技术
    131: 1,  # 2.1 后端开发技术
    132: 2,  # 2.1.1 Spring Boot 框架
    134: 2,  # 2.1.2 MyBatis-Plus 持久层框架
    136: 2,  # 2.1.3 MySQL 与 Redis
    138: 2,  # 2.1.4 WebSocket 即时通讯
    140: 2,  # 2.1.5 JWT 认证机制
    142: 2,  # 2.1.6 定时任务
    144: 1,  # 2.2 前端开发技术
    145: 2,  # 2.2.1 uni-app 跨端框架
    147: 2,  # 2.2.2 Vue 3 组合式 API 与 Pinia
    149: 2,  # 2.2.3 Element Plus 与 ECharts
    # Chapter 3
    152: 0,  # 三、系统分析
    153: 1,  # 3.1 可行性分析
    154: 2,  # 3.1.1 技术可行性
    156: 2,  # 3.1.2 操作可行性
    158: 1,  # 3.2 系统需求分析
    159: 2,  # 3.2.1 功能性需求
    162: 2,  # 3.2.2 非功能性需求
    164: 1,  # 3.3 系统用例分析
    165: 2,  # 3.3.1 用户端用例
    167: 2,  # 3.3.2 管理端用例
    169: 2,  # 3.3.3 核心业务流程
    # Chapter 4
    172: 0,  # 四、系统设计
    173: 1,  # 4.1 系统架构设计
    174: 2,  # 4.1.1 总体架构
    176: 2,  # 4.1.2 模块划分
    178: 1,  # 4.2 数据库设计
    179: 2,  # 4.2.1 核心实体设计
    181: 2,  # 4.2.2 关键数据表结构
    199: 2,  # 4.2.3 Redis 缓存设计
    204: 1,  # 4.3 接口设计
    205: 2,  # 4.3.1 统一响应格式
    207: 2,  # 4.3.2 小程序端接口概览
    209: 2,  # 4.3.3 管理端接口概览
    # Chapter 5
    212: 0,  # 五、系统实现
    213: 1,  # 5.1 后端核心功能实现
    214: 2,  # 5.1.1 JWT 鉴权拦截器
    216: 2,  # 5.1.2 WebSocket 即时通信
    218: 2,  # 5.1.3 消息协议设计
    220: 2,  # 5.1.4 商品状态管理与订单流转
    222: 2,  # 5.1.5 文件上传与数据统计
    224: 1,  # 5.2 前端核心功能实现
    225: 2,  # 5.2.1 微信小程序页面实现
    228: 2,  # 5.2.2 全局状态管理
    230: 2,  # 5.2.3 管理后台页面实现
    232: 2,  # 5.3 关键界面展示
    # Chapter 6
    235: 0,  # 六、系统测试与总结
    236: 1,  # 6.1 系统测试
    237: 2,  # 6.1.1 功能测试
    240: 2,  # 6.1.2 兼容性测试
    242: 1,  # 6.2 全文总结
    245: 1,  # 6.3 不足与展望
}

for idx, level in heading_fixes.items():
    if idx < len(doc.paragraphs):
        set_outline_level(doc.paragraphs[idx], level)

print(f'Set outlineLvl for {len(heading_fixes)} headings')

# Step 3: Now insert the 6 images (same as before)
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

IMG_DIR = r'D:\Desktop\笔记\使用照片'

insertions = [
    (220, '订单状态机流转图.png',       '图5-3  订单状态机流转图'),
    (216, 'WebSocket 即时通讯架构图.png', '图5-2  WebSocket即时通讯架构图'),
    (179, '数据库简化 ER 关系图.png',   '图4-2  数据库ER关系图'),
    (174, '技术架构分层图.png',         '图4-1  系统技术架构分层图'),
    (169, '商品发布审核流程图.png',     '图3-2  商品发布审核流程图'),
    (165, '校园认证流程图.png',         '图3-1  校园认证流程图'),
]

def insert_image_with_caption(doc, after_idx, img_path, caption_text):
    p_blank1 = doc.add_paragraph()
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(14))
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.size = Pt(10.5)
    run_cap.font.name = '宋体'
    run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p_blank2 = doc.add_paragraph()

    target = doc.paragraphs[after_idx]._element
    body = target.getparent()
    all_paras = body.findall(qn('w:p'))
    new_elems = [all_paras[-4], all_paras[-3], all_paras[-2], all_paras[-1]]
    for elem in new_elems:
        body.remove(elem)
    insert_after = target
    for elem in new_elems:
        insert_after.addnext(elem)
        insert_after = elem

insertions_sorted = sorted(insertions, key=lambda x: x[0], reverse=True)
for after_idx, img_file, caption in insertions_sorted:
    img_path = f'{IMG_DIR}\\{img_file}'
    print(f'Inserting {img_file} after P{after_idx}...')
    insert_image_with_caption(doc, after_idx, img_path, caption)

# Step 4: Save with updateFields
doc.save(OUT)
print(f'\nSaved to: {OUT}')

# Now set updateFields in the saved docx
import zipfile, os, shutil, tempfile
from lxml import etree

tmp_dir = tempfile.mkdtemp()
with zipfile.ZipFile(OUT, 'r') as z:
    z.extractall(tmp_dir)

settings_path = os.path.join(tmp_dir, 'word', 'settings.xml')
if os.path.exists(settings_path):
    tree = etree.parse(settings_path)
    root = tree.getroot()
    update_fields = root.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
    if update_fields is not None:
        update_fields.set('{http://www.w3.org/XML/1998/namespace}val', 'true')
    else:
        uf = etree.SubElement(root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields')
        uf.set('{http://www.w3.org/XML/1998/namespace}val', 'true')
    tree.write(settings_path, xml_declaration=True, encoding='UTF-8', standalone=True)

os.remove(OUT)
with zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for dirpath, dirnames, filenames in os.walk(tmp_dir):
        for f in filenames:
            full_path = os.path.join(dirpath, f)
            arc_path = os.path.relpath(full_path, tmp_dir)
            z.write(full_path, arc_path)
shutil.rmtree(tmp_dir)

print('TOC will auto-update on open in Word.')
