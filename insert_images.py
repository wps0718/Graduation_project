# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import io

IMG_DIR = r'G:\Code\PPT\ppt-master-main\projects\graduation_defense_ppt169_20260511\images'
SRC_DOC = r'G:\Code\Graduation_project\毕业设计论文_精进版.docx'
OUT_DOC = r'G:\Code\Graduation_project\毕业设计论文_最终版.docx'

doc = Document(SRC_DOC)

# Image insertion plan: (after_paragraph_index, image_filename, caption)
# We insert in reverse order (highest index first) to avoid index shifting
insertions = [
    (97,  'WebSocket 即时通讯架构图.png', '图5-2  WebSocket即时通讯架构图'),
    (101, '订单状态机流转图.png',         '图5-3  订单状态机流转图'),
    (77,  '数据库简化 ER 关系图.png',     '图4-1  数据库ER关系图'),
    (72,  '技术架构分层图.png',           '图4-2  系统技术架构分层图'),
    (67,  '商品发布审核流程图.png',       '图3-1  商品发布审核流程图'),
    (63,  '校园认证流程图.png',           '图3-2  校园认证流程图'),
]

def add_image_paragraph(doc, after_index, img_path, caption_text):
    """Insert an image + caption after the given paragraph index."""
    # We need to insert after the paragraph at after_index
    # In python-docx, we can't easily insert at arbitrary positions,
    # so we append and then reorder, OR we manipulate the XML directly.

    # Strategy: use XML manipulation to insert new paragraphs after the target
    from lxml import etree

    # Create the image paragraph
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(14))  # ~14cm wide, fits A4

    # Create the caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = p_cap.add_run(caption_text)
    run_cap.font.size = Pt(10.5)
    run_cap.font.name = '宋体'
    run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Add a blank paragraph for spacing
    p_blank = doc.add_paragraph()

    # Now move these 3 paragraphs from the end to after our target paragraph
    target_para = doc.paragraphs[after_index]
    target_element = target_para._element

    # Get the 3 newly added paragraph elements (at the end)
    body = target_element.getparent()
    all_paras = body.findall(qn('w:p'))

    # The last 3 elements are our new ones
    new_elements = [all_paras[-1], all_paras[-2], all_paras[-3]]  # blank, caption, image (reversed)

    # Insert them after target_element, in correct order: image, caption, blank
    insert_after = target_element
    for elem in reversed(new_elements):  # image first, then caption, then blank
        body.remove(elem)
        insert_after.addnext(elem)
        insert_after = elem

# Sort by index descending so we insert from bottom to top
insertions_sorted = sorted(insertions, key=lambda x: x[0], reverse=True)

for after_idx, img_file, caption in insertions_sorted:
    img_path = f'{IMG_DIR}\\{img_file}'
    print(f'Inserting {img_file} after paragraph {after_idx}...')
    add_image_paragraph(doc, after_idx, img_path, caption)

doc.save(OUT_DOC)
print(f'\nSaved to: {OUT_DOC}')

# Count final characters
import re
total = 0
chinese = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        total += len(text)
        chinese += len(re.findall(r'[一-鿿]', text))
print(f'Total characters: {total}')
print(f'Chinese characters: {chinese}')

# Count images
import zipfile
with zipfile.ZipFile(OUT_DOC) as z:
    imgs = [f for f in z.namelist() if f.startswith('word/media/')]
    print(f'Images in document: {len(imgs)}')
