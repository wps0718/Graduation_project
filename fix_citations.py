# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

SRC = r'D:\Desktop\毕业设计(论文）最最最总.docx'
OUT = r'D:\Desktop\毕业设计(论文）最最最总_修正版.docx'

doc = Document(SRC)

# ========== Step 1: Fix in-text citations ==========
# Old → New mapping for paragraphs where the citation number is wrong
in_text_fixes = {
    137: ('[4]', '[9]'),   # MySQL/Redis → should cite Redis in Action [9]
    139: ('[5]', '[4]'),   # WebSocket → should cite 即时通讯论文 [4]
    141: ('[6]', '[5]'),   # JWT → should cite RFC 7519 [5]
    146: ('[7]', '[6]'),   # uni-app → should cite DCloud uni-app [6]
    148: ('[8]', '[7]'),   # Vue 3 → should cite Vue 3 docs [7]
    150: ('[9]', '[8]'),   # ECharts → should cite ECharts论文 [8]
    200: ('[7]', '[9]'),   # Redis → should cite Redis in Action [9]
}

for idx, (old_cite, new_cite) in in_text_fixes.items():
    para = doc.paragraphs[idx]
    old_text = para.text
    if old_cite in old_text:
        # Preserve bold/italic by manipulating runs
        for run in para.runs:
            if old_cite in run.text:
                run.text = run.text.replace(old_cite, new_cite)
                break
        else:
            # Citation might span multiple runs, fall back to full text replacement
            was_bold = any(r.bold for r in para.runs)
            for run in para.runs:
                run.text = ''
            para.runs[0].text = old_text.replace(old_cite, new_cite)
            if was_bold:
                para.runs[0].bold = True
        print(f'P{idx}: {old_cite} -> {new_cite}')
    else:
        print(f'P{idx}: WARNING - {old_cite} not found in text')

# ========== Step 2: Fix reference list (P255-P262) ==========
# New reference list in correct order:
new_refs = {
    255: '[1] 郭林盛,张晓艳,赵妍,李玥. 基于微信云开发实现校园二手交易的平台设计[J]. 现代计算机, 2024, 30(22): 185-190.',
    256: '[2] Walls C. Spring Boot in Action[M]. Manning Publications, 2016.',
    257: '[3] MyBatis-Plus官方团队. MyBatis-Plus官方文档[EB/OL]. https://baomidou.com, 2024.',
    258: '[4] 杜瑞庆,李一诺. 基于Java语言的即时通讯系统设计与实现[J]. 电脑知识与技术, 2022, 18(31): 29-32.',
    259: '[5] Jones M, Bradley J, Sakimura N. JSON Web Token (JWT)[S]. RFC 7519, IETF, 2015.',
    260: '[6] DCloud. uni-app跨端开发官方文档[EB/OL]. https://uniapp.dcloud.net.cn, 2024.',
    261: '[7] You Y. Vue 3 Composition API Documentation[EB/OL]. https://vuejs.org, 2024.',
    262: '[8] 刘梦,张浩洋,唐松强. 基于ECharts的可视化技术在数据管理平台中的应用[J]. 现代计算机, 2024, 30(11): 105-109.',
}

for idx, new_text in new_refs.items():
    para = doc.paragraphs[idx]
    was_bold = any(r.bold for r in para.runs)
    for run in para.runs:
        run.text = ''
    para.runs[0].text = new_text
    if was_bold:
        para.runs[0].bold = True
    print(f'P{idx}: Updated reference')

# ========== Step 3: Add [9] Redis in Action as a new paragraph after P262 ==========
# We need to add a new paragraph for [9]. Since we can't easily insert,
# we'll append at the end near the reference section.
# Check if there's a blank paragraph after P262 we can use, or add one.

# Find the paragraph after P262
added = False
for i in range(263, min(270, len(doc.paragraphs))):
    para = doc.paragraphs[i]
    text = para.text.strip()
    # If there's an empty paragraph or a formatting instruction page, we can use it
    if text == '' or text.startswith('格式要求') or text.startswith('附录'):
        # Insert before this paragraph
        # We'll set the text of the blank line before the formatting page
        prev_idx = i - 1
        if prev_idx > 262:
            para_prev = doc.paragraphs[prev_idx]
            if para_prev.text.strip() == '':
                para_prev.add_run('[9] Carlson J. Redis in Action[M]. Manning Publications, 2016.')
                added = True
                print(f'P{prev_idx}: Added [9] Redis in Action')
                break

if not added:
    # Try to add after P262 directly
    para = doc.paragraphs[262]
    # Check next paragraph
    if 263 < len(doc.paragraphs):
        next_para = doc.paragraphs[263]
        if next_para.text.strip() == '':
            next_para.add_run('[9] Carlson J. Redis in Action[M]. Manning Publications, 2016.')
            added = True
            print(f'P263: Added [9] Redis in Action')

if not added:
    print('WARNING: Could not find a place to add [9]. Will need manual insertion.')

# Save
doc.save(OUT)
print(f'\nSaved to: {OUT}')
